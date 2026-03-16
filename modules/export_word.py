from docx import Document

def export_word(equation, solution, explanation):

    file_path = "outputs/solution.docx"

    doc = Document()

    doc.add_heading("AI Math Tutor Solution", level=1)

    doc.add_paragraph(f"Equation: {equation}")

    doc.add_paragraph(f"Solution: x = {solution}")

    doc.add_heading("Explanation", level=2)

    doc.add_paragraph(explanation)

    doc.save(file_path)

    return file_path